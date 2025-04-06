from flask import Flask, request, send_file, jsonify
import pandas as pd
import os
import random
import boto3
from flask import Flask, request, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from groq import Groq
import os
import base64
from flask import Flask, request, jsonify
import tempfile
from werkzeug.utils import secure_filename
from flask import Flask, request, send_file, jsonify
from moviepy import VideoFileClip
from pydub import AudioSegment
import os
import uuid
from PIL import Image
import io
from flask import Flask, request, send_file, jsonify
from pdf2docx import Converter
import os
import tempfile
from flask import Flask, request, jsonify, send_file
from docx2pdf import convert
import tempfile
import uuid
from flask import Flask, request, jsonify, send_file, after_this_request
from docx2pdf import convert
import tempfile
import uuid
import shutil
import requests
from flask import Flask, request, jsonify, send_file, after_this_request,send_from_directory
import PyPDF2
from pypdf import PdfWriter, PdfReader
from docx import Document
from flask_cors import CORS
from PIL import Image
import tempfile
from flask import Flask, request, jsonify
import psycopg2
import psycopg2.extras
import re
import google.generativeai as genai
import os
from urllib.parse import urlparse
from dotenv import load_dotenv
from flask import Flask, request, jsonify, send_file
from PIL import Image, ImageDraw
import extcolors
import math
import io
from flask import Flask, request, jsonify
import base64
import os
from mistralai import Mistral
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import matplotlib.pyplot as plt
from matplotlib import gridspec
import base64
import random
import uuid
from flask import Flask, request, jsonify
from dotenv import load_dotenv
import requests
from openai import OpenAI
from dotenv import load_dotenv
from flask import Flask, request, jsonify
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import os
UPLOAD_FOLDER = "uploads"
PROCESSED_FOLDER = "processed"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
load_dotenv()
app = Flask(__name__)
CORS(app)
api_key = os.getenv('CURRENCY_FREAKS')
# Ensure temporary directory exists
UPLOAD_FOLDER = os.path.join(os.getcwd(), "temp_uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
def ensure_uploads_dir():
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
def get_pdf_text(pdf_files):
    text = ""
    for pdf in pdf_files:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Function to split text into chunks
def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    return text_splitter.split_text(text)

# Function to create FAISS vector store
def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

# Function to create the conversational chain
def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context. 
    If the answer is not in the context, respond with: "Answer is not available in the context."
    
    Context:\n {context}\n
    Question:\n{question}\n
    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", temperature=0.3)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable not set")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel('gemini-1.5-pro')

# Function to parse PostgreSQL connection URL
def parse_db_url(db_url):
    result = urlparse(db_url)
    return {
        "dbname": result.path[1:],
        "user": result.username,
        "password": result.password,
        "host": result.hostname,
        "port": result.port or "5432",
        "sslmode": "require"
    }

# Function to connect to the database with connection details
def get_db_connection(db_config):
    return psycopg2.connect(**db_config)

# Get all table names
@app.route("/tables", methods=["GET"])
def get_table_names():
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
        
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor()
        cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'")
        tables = [row[0] for row in cur.fetchall()]
        cur.close()
        conn.close()
        return jsonify({"tables": tables})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Initialize Groq client with API key
client_groq = Groq(api_key=GROQ_API_KEY)
languages = {
    'en': 'English',
    'hi': 'Hindi',
    'gu': 'Gujarati',
    'ma': 'Marathi'
}

llm = ChatGroq(
    model="llama-3.3-70b-versatile",
    temperature=0.7
)

# Updated system prompt with language instruction
def get_system_prompt(lang='en'):
    base_prompt = """
# Coding Assistant Role
You are an expert AI coding assistant specialized in helping users generate, debug, and optimize code. Your purpose is to provide clear, accurate, and well-explained solutions to coding challenges across multiple programming languages and paradigms.

# Expertise Areas
- Software design and architecture
- Algorithm implementation and optimization
- Language-specific best practices
- Performance tuning and efficiency improvements
- Debugging and error resolution
- Testing and test-driven development
- Refactoring and code maintenance
- Security considerations and vulnerabilities

# Response Structure
For each user query, follow this structure:

## Understanding the Problem
1. Restate the user's goal or issue to confirm understanding
2. Identify any unstated assumptions or edge cases to consider
3. Note any potential constraints or performance requirements

## Solution Approach
1. Outline your overall approach to solving the problem
2. Explain why this approach is appropriate for this specific case
3. Mention alternative approaches if relevant, and why you chose this one

## Code Implementation
1. Provide complete, runnable code that solves the problem
2. Include appropriate error handling, input validation, and edge case management
3. Structure the code with proper organization, naming conventions, and documentation

## Code Explanation
1. Walk through the implementation line-by-line or section-by-section
2. Explain the purpose and function of critical components
3. Highlight any non-obvious or complex logic

## Optimization Insights
1. Identify performance characteristics (time/space complexity)
2. Suggest optimizations if the original solution can be improved
3. Explain the tradeoffs involved in any optimization decisions

## Testing Considerations
1. Provide example test cases covering typical usage and edge cases
2. Suggest testing methodologies appropriate for this code
3. Discuss potential failure modes and how to handle them

## Best Practices
1. Highlight language-specific idioms and conventions used
2. Note any design patterns or architectural principles applied
3. Reference relevant documentation, libraries, or resources for further learning

# Programming Languages
Demonstrate expertise across popular languages including but not limited to:
- Python, JavaScript/TypeScript, Java, C#, C/C++, Go, Rust, PHP, Ruby, Swift, Kotlin
- SQL and database query languages
- Shell scripting (Bash, PowerShell)
- Web technologies (HTML, CSS, various JS frameworks)

# Special Instructions
- When debugging, analyze the error systematically and provide clear explanations for the cause
- When optimizing, consider both algorithmic improvements and language-specific optimizations
- When suggesting refactoring, explain the benefits in terms of readability, maintainability, and performance
- Include meaningful comments in code examples to enhance understanding
- Support answers with relevant computer science principles and design philosophies
- Adapt your technical depth to match the apparent expertise level of the user
- Emphasize not just what to do, but why to do it that way

Question: {question}

Answer:
"""
    lang_instruction = f"\nPlease provide all responses in {languages.get(lang, 'English')}."
    return base_prompt + lang_instruction

@app.route('/code_assistant', methods=['POST'])
def chat():
    data = request.json
    user_query = data.get('query')
    lang = data.get('language', 'en')  # Default to English if no language specified
    
    if not user_query:
        return jsonify({"error": "Query is required"}), 400
    
    if lang not in languages:
        return jsonify({"error": f"Unsupported language code. Supported codes are: {', '.join(languages.keys())}"}), 400
    
    # Combine system prompt with user query
    messages = [
        {"role": "system", "content": get_system_prompt(lang)},
        {"role": "user", "content": user_query}
    ]
    
    try:
        # Get response from LLM
        response = llm.invoke(messages)
        return jsonify({
            "response": response.content,
            "language": languages[lang],
            "status": "success"
        })
    except Exception as e:
        return jsonify({
            "error": str(e),
            "status": "error"
        }), 500
# Get schema information for all tables
api_key_mistral = os.environ.get("MISTRAL_API_KEY")
if not api_key_mistral:
    raise ValueError("MISTRAL_API_KEY environment variable is not set.")

# Initialize Mistral client_mistral
client_mistral = Mistral(api_key=api_key_mistral)

def encode_image_file(file_storage):
    """Encode uploaded image file to base64."""
    try:
        return base64.b64encode(file_storage.read()).decode("utf-8")
    except Exception as e:
        print(f"Encoding error: {e}")
        return None

@app.route("/ocr", methods=["POST"])
def ocr_from_image():
    if "image" not in request.files:
        return jsonify({"error": "No image file provided"}), 400

    image_file = request.files["image"]

    if image_file.filename == "":
        return jsonify({"error": "Empty filename"}), 400

    try:
        # Secure the filename and encode the image
        filename = secure_filename(image_file.filename)
        base64_image = encode_image_file(image_file)

        if not base64_image:
            return jsonify({"error": "Failed to encode image"}), 500

        # Call Mistral OCR
        ocr_response = client_mistral.ocr.process(
            model="mistral-ocr-latest",
            document={
                "type": "image_url",
                "image_url": f"data:image/jpeg;base64,{base64_image}"
            }
        )
        text = "\n".join([page.markdown for page in ocr_response.pages])
        return jsonify({"text": text})

    except Exception as e:
        print(f"OCR error: {e}")
        return jsonify({"error": str(e)}), 500
@app.route("/schema", methods=["GET"])
def get_schema():
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
            
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Get all tables
        cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'")
        tables = [row[0] for row in cur.fetchall()]

        schema_info = {}
        for table in tables:
            # Get column information
            cur.execute(f"SELECT column_name, data_type, is_nullable FROM information_schema.columns WHERE table_name = %s", (table,))
            columns = [{"name": row["column_name"], "type": row["data_type"], "nullable": row["is_nullable"] == "YES"} for row in cur.fetchall()]
            
            # Get primary key information
            cur.execute("""
                SELECT kcu.column_name 
                FROM information_schema.table_constraints tc
                JOIN information_schema.key_column_usage kcu
                    ON tc.constraint_name = kcu.constraint_name
                WHERE tc.constraint_type = 'PRIMARY KEY' 
                AND tc.table_name = %s
            """, (table,))
            primary_keys = [row[0] for row in cur.fetchall()]
            
            # Get foreign key information
            cur.execute("""
                SELECT
                    kcu.column_name,
                    ccu.table_name AS foreign_table_name,
                    ccu.column_name AS foreign_column_name
                FROM information_schema.table_constraints AS tc
                JOIN information_schema.key_column_usage AS kcu
                    ON tc.constraint_name = kcu.constraint_name
                JOIN information_schema.constraint_column_usage AS ccu
                    ON ccu.constraint_name = tc.constraint_name
                WHERE tc.constraint_type = 'FOREIGN KEY' 
                AND tc.table_name = %s
            """, (table,))
            foreign_keys = [{"column": row[0], "references_table": row[1], "references_column": row[2]} for row in cur.fetchall()]
            
            schema_info[table] = {
                "columns": columns,
                "primary_keys": primary_keys,
                "foreign_keys": foreign_keys
            }

        cur.close()
        conn.close()
        return jsonify(schema_info)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/convert_video', methods=['POST'])
def convert_video():
    file = request.files.get('file')
    output_format = request.form.get('format')
    if not file or not output_format:
        return jsonify({'error': 'Missing file or output format'}), 400
    
    filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[-1]
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)

    clip = VideoFileClip(input_path)
    clip.write_videofile(output_path)

    return send_file(output_path, as_attachment=True)

import os
import uuid
import asyncio
from flask import Flask, request, jsonify, send_file, abort
from pydub import AudioSegment
from concurrent.futures import ThreadPoolExecutor
# --- Trim Video ---
@app.route('/trim_video', methods=['POST'])
def trim_video():
    file = request.files.get('file')
    start_time = float(request.form.get('start', 0))
    end_time = float(request.form.get('end', 0))

    if not file:
        return jsonify({'error': 'Missing file'}), 400

    filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[-1]
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    output_filename = f"{uuid.uuid4()}.mp4"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)

    clip = VideoFileClip(input_path).subclipped(start_time, end_time)
    clip.write_videofile(output_path)

    return send_file(output_path, as_attachment=True)


# --- Audio Conversion ---
@app.route('/convert_audio', methods=['POST'])
def convert_audio():
    file = request.files.get('file')
    output_format = request.form.get('format')

    if not file or not output_format:
        return jsonify({'error': 'Missing file or output format'}), 400

    filename = str(uuid.uuid4()) + os.path.splitext(file.filename)[-1]
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    output_filename = f"{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)

    audio = AudioSegment.from_file(input_path)
    audio.export(output_path, format=output_format)

    return send_file(output_path, as_attachment=True)

executor = ThreadPoolExecutor()

def trim_and_export(input_path, output_path, start, end):
    audio = AudioSegment.from_file(input_path)
    trimmed = audio[start:end]
    trimmed.export(output_path, format="mp3")
    return output_path

@app.route('/trim_audio', methods=['POST'])
async def trim_audio():
    file = request.files.get('file')
    start_time = int(request.form.get('start', 0))
    end_time = int(request.form.get('end', 0))

    if not file:
        return jsonify({'error': 'Missing file'}), 400

    filename = f"{uuid.uuid4()}{os.path.splitext(file.filename)[-1]}"
    input_path = os.path.join(UPLOAD_FOLDER, filename)
    file.save(input_path)

    output_filename = f"{uuid.uuid4()}.mp3"
    output_path = os.path.join(PROCESSED_FOLDER, output_filename)

    try:
        # Run blocking audio processing in a separate thread
        await asyncio.get_event_loop().run_in_executor(
            executor,
            trim_and_export,
            input_path,
            output_path,
            start_time,
            end_time
        )

        if not os.path.exists(output_path):
            raise FileNotFoundError(f"Export failed, file not found at: {output_path}")

        return send_file(output_path, as_attachment=True)

    except Exception as e:
        print(f"[ERROR] {str(e)}")
        return jsonify({'error': 'Failed to process audio', 'details': str(e)}), 500
import time
@app.route('/api/test', methods=['POST'])
def api_tester():
    try:
        input_data = request.get_json()

        url = input_data.get("url")
        method = input_data.get("method", "GET").upper()
        headers = input_data.get("headers", {})
        data = input_data.get("data", None)
        token = input_data.get("token", None)

        # Add Bearer token if provided
        if token:
            headers["Authorization"] = f"Bearer {token}"

        # Measure latency
        start_time = time.time()
        response = requests.request(method=method, url=url, headers=headers, json=data)
        latency = (time.time() - start_time) * 1000  # ms

        return jsonify({
            "status_code": response.status_code,
            "response": response.json() if "application/json" in response.headers.get("Content-Type", "") else response.text,
            "latency_ms": round(latency, 2)
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
# Get all records from a table
@app.route("/<string:table_name>", methods=["GET"])
def get_all_records(table_name):
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
            
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute(f"SELECT * FROM {table_name}")
        data = [dict(row) for row in cur.fetchall()]
        cur.close()
        conn.close()
        return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
# New endpoint to generate SQL based on natural language
@app.route("/generate-sql", methods=["POST"])
def generate_sql():
    try:
        data = request.json
        if not data or 'db_url' not in data or 'query' not in data:
            return jsonify({"error": "Missing required parameters: db_url and query"}), 400
        
        db_url = data['db_url']
        natural_language_query = data['query']
        
        # Get schema information
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Get all tables
        cur.execute("SELECT table_name FROM information_schema.tables WHERE table_schema = 'public'")
        tables = [row[0] for row in cur.fetchall()]

        schema_info = {}
        for table in tables:
            # Get column information
            cur.execute(f"SELECT column_name, data_type, is_nullable FROM information_schema.columns WHERE table_name = %s", (table,))
            columns = [{"name": row["column_name"], "type": row["data_type"], "nullable": row["is_nullable"] == "YES"} for row in cur.fetchall()]
            
            # Get primary key information
            cur.execute("""
                SELECT kcu.column_name 
                FROM information_schema.table_constraints tc
                JOIN information_schema.key_column_usage kcu
                    ON tc.constraint_name = kcu.constraint_name
                WHERE tc.constraint_type = 'PRIMARY KEY' 
                AND tc.table_name = %s
            """, (table,))
            primary_keys = [row[0] for row in cur.fetchall()]
            
            # Get foreign key information
            cur.execute("""
                SELECT
                    kcu.column_name,
                    ccu.table_name AS foreign_table_name,
                    ccu.column_name AS foreign_column_name
                FROM information_schema.table_constraints AS tc
                JOIN information_schema.key_column_usage AS kcu
                    ON tc.constraint_name = kcu.constraint_name
                JOIN information_schema.constraint_column_usage AS ccu
                    ON ccu.constraint_name = tc.constraint_name
                WHERE tc.constraint_type = 'FOREIGN KEY' 
                AND tc.table_name = %s
            """, (table,))
            foreign_keys = [{"column": row[0], "references_table": row[1], "references_column": row[2]} for row in cur.fetchall()]
            
            schema_info[table] = {
                "columns": columns,
                "primary_keys": primary_keys,
                "foreign_keys": foreign_keys
            }

        cur.close()
        conn.close()
        
        # Format schema information for the prompt
        schema_text = "Database Schema:\n"
        for table, info in schema_info.items():
            schema_text += f"Table: {table}\n"
            schema_text += "  Columns:\n"
            for col in info["columns"]:
                nullable = "NULL" if col["nullable"] else "NOT NULL"
                schema_text += f"    - {col['name']} ({col['type']}, {nullable})\n"
            
            if info["primary_keys"]:
                schema_text += "  Primary Keys:\n"
                for pk in info["primary_keys"]:
                    schema_text += f"    - {pk}\n"
            
            if info["foreign_keys"]:
                schema_text += "  Foreign Keys:\n"
                for fk in info["foreign_keys"]:
                    schema_text += f"    - {fk['column']} references {fk['references_table']}({fk['references_column']})\n"
            
            schema_text += "\n"
        
        # Create prompt for Gemini
        prompt = f"""
You are an expert SQL query generator. Based on the following database schema, generate a PostgreSQL SQL query that answers the user's question.
Follow these rules:
1. Only generate the raw SQL query, nothing else
2. Use proper table and column names as shown in the schema
3. Include appropriate JOINs if the query requires data from multiple tables
4. Return the query without backticks, markdown formatting, or any explanations
5. Be mindful of the relationships between tables

{schema_text}

User Question: {natural_language_query}

SQL Query:
"""
        
        # Get response from Gemini
        response = model.generate_content(prompt)
        
        # Clean the response to extract only the SQL query
        sql_query = response.text.strip()
        
        # Remove any markdown code formatting if present
        sql_query = re.sub(r'^```sql\s*', '', sql_query)
        sql_query = re.sub(r'\s*```$', '', sql_query)
        
        # Try to execute the query to validate it
        try:
            conn = get_db_connection(db_config)
            cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
            cur.execute(sql_query)
            
            # If SELECT query, return the results
            if sql_query.strip().lower().startswith("select"):
                results = [dict(row) for row in cur.fetchall()]
                conn.commit()
                cur.close()
                conn.close()
                return jsonify({
                    "sql": sql_query,
                    "results": results,
                    "rows_affected": len(results)
                })
            else:
            # For non-SELECT queries, capture the affected rows
                rows_affected = cur.rowcount
                
                # Try to determine which table was modified
                table_name = None
                # Extract table name from the query (basic parsing)
                query_lower = sql_query.lower()
                if "update " in query_lower:
                    # Extract table name from UPDATE query
                    match = re.search(r'update\s+([^\s]+)', query_lower)
                    if match:
                        table_name = match.group(1)
                elif "insert into " in query_lower:
                    # Extract table name from INSERT query
                    match = re.search(r'insert\s+into\s+([^\s\(]+)', query_lower)
                    if match:
                        table_name = match.group(1)
                elif "delete from " in query_lower:
                    # Extract table name from DELETE query
                    match = re.search(r'delete\s+from\s+([^\s]+)', query_lower)
                    if match:
                        table_name = match.group(1)
                
                # If we identified a table, return its contents
                if table_name:
                    # Remove any schema prefix from the table name if present
                    if "." in table_name:
                        table_name = table_name.split(".")[-1]
                    # Also remove any quotes from the table name
                    table_name = table_name.strip('"\'')
                    
                    try:
                        # Fetch the updated table contents
                        select_query = f'SELECT * FROM "{table_name}"'
                        cur.execute(select_query)
                        table_data = [dict(row) for row in cur.fetchall()]
                        
                        # Commit changes and close connection
                        conn.commit()
                        cur.close()
                        conn.close()
                        
                        return jsonify({
                            "sql": sql_query,
                            "rows_affected": rows_affected,
                            "table_affected": table_name,
                            "table_data": table_data
                        })
                    except Exception as e:
                        # If we can't get the table data, just return the original response
                                pass
                        
                        # Default case if we couldn't get table data
                conn.commit()
                cur.close()
                conn.close()
                return jsonify({
                            "sql": sql_query,
                            "rows_affected": rows_affected
                        })
        except Exception as e:
                        # Return the SQL with error if execution fails
                        return jsonify({
                            "sql": sql_query,
                            "error": str(e),
                            "note": "SQL generation succeeded, but execution failed. You may need to modify the query."
                        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Create a new record
@app.route("/<string:table_name>", methods=["POST"])
def create_record(table_name):
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
            
        data = request.json
        columns = ", ".join(data.keys())
        values = ", ".join(["%s"] * len(data))
        query = f"INSERT INTO {table_name} ({columns}) VALUES ({values}) RETURNING *"

        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cur.execute(query, tuple(data.values()))
        result = dict(cur.fetchone())
        conn.commit()
        cur.close()
        conn.close()

        return jsonify(result), 201
    except Exception as e:
        return jsonify({"error": str(e)}), 400

# Update a record by primary key
@app.route("/<string:table_name>/<primary_key_value>", methods=["PUT"])
def update_record(table_name, primary_key_value):
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
            
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Get primary key
        cur.execute(f"SELECT column_name FROM information_schema.key_column_usage WHERE table_name = %s", (table_name,))
        primary_key = cur.fetchone()
        if not primary_key:
            return jsonify({"error": "Primary key not found"}), 400
        primary_key = primary_key[0]

        data = request.json
        update_query = ", ".join([f"{key} = %s" for key in data.keys()])
        query = f"UPDATE {table_name} SET {update_query} WHERE {primary_key} = %s RETURNING *"

        cur.execute(query, tuple(data.values()) + (primary_key_value,))
        result = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        if result:
            return jsonify(dict(result))
        return jsonify({"message": "Record not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 400

# Delete a record by primary key
@app.route("/<string:table_name>/<primary_key_value>", methods=["DELETE"])
def delete_record(table_name, primary_key_value):
    try:
        db_url = request.args.get('db_url')
        if not db_url:
            return jsonify({"error": "Missing db_url parameter"}), 400
            
        db_config = parse_db_url(db_url)
        conn = get_db_connection(db_config)
        cur = conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        
        # Get primary key
        cur.execute(f"SELECT column_name FROM information_schema.key_column_usage WHERE table_name = %s", (table_name,))
        primary_key = cur.fetchone()
        if not primary_key:
            return jsonify({"error": "Primary key not found"}), 400
        primary_key = primary_key[0]

        query = f"DELETE FROM {table_name} WHERE {primary_key} = %s RETURNING *"
        cur.execute(query, (primary_key_value,))
        result = cur.fetchone()
        conn.commit()
        cur.close()
        conn.close()

        if result:
            return jsonify({"message": "Record deleted"})
        return jsonify({"message": "Record not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 400
# Endpoint to process PDFs
@app.route("/process_pdfs", methods=["POST"])
def process_pdfs():
    try:
        if "files" not in request.files:
            return jsonify({"error": "No file part"}), 400
        
        pdf_files = request.files.getlist("files")
        raw_text = get_pdf_text(pdf_files)
        text_chunks = get_text_chunks(raw_text)
        get_vector_store(text_chunks)

        return jsonify({"message": "Processing complete"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route('/summarize', methods=['POST'])
def summarize_text():
    # Get request data
    data = request.json
    
    if not data or 'text' not in data:
        return jsonify({"error": "No text provided"}), 400
    
    text = data['text']
    output_length = data.get('output_length', 100)  # Default to 100 if not specified
    
    try:
        # Initialize the Gemini model
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # Create prompt with instruction for summary length
        prompt = f"""
        Summarize the following text in approximately {output_length} words:
        
        {text}
        """
        
        # Generate summary
        response = model.generate_content(prompt)
        summary = response.text
        
        return jsonify({
            "summary": summary,
            "requested_length": output_length
        })
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500
ALLOWED_VOICES = {
    "alloy", "ash", "ballad", "coral", "echo",
    "fable", "onyx", "nova", "sage", "shimmer"
}

@app.route("/tts", methods=["POST"])
def generate_speech():
    data = request.json

    if not data or "text" not in data or "voice" not in data:
        return jsonify({"error": "Missing 'text' or 'voice' in request"}), 400

    text = data["text"]
    voice = data["voice"]

    if voice not in ALLOWED_VOICES:
        return jsonify({"error": f"Invalid voice '{voice}'. Choose from {list(ALLOWED_VOICES)}"}), 400

    try:    
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio_file:
            with client.audio.speech.with_streaming_response.create(
                model="tts-1",  # Or "tts-1-hd" if preferred
                voice=voice,
                input=text
            ) as response:
                response.stream_to_file(temp_audio_file.name)

            return send_file(temp_audio_file.name, mimetype="audio/mpeg", as_attachment=True, download_name="speech.mp3")

    except Exception as e:
        return jsonify({"error": str(e)}), 500
# Endpoint to handle user questions
s3=boto3.client('s3')
S3_BUCKET = os.environ.get("S3_BUCKET", "testusebucket123")
@app.route("/upload", methods=["POST"])
def upload_and_get_presigned_url():
    # 1) Validate file
    if "file" not in request.files:
        return jsonify({"error": "No file part in the request"}), 400
    file_obj = request.files["file"]
    if file_obj.filename == "":
        return jsonify({"error": "No file selected"}), 400

    # 2) Validate expiration
    expiration = request.form.get("expiration", type=int)
    if expiration is None or expiration <= 0:
        return jsonify({"error": "Invalid expiration; must be a positive integer"}), 400

    # 3) Generate a unique object key
    object_key = f"obj{random.randint(1, 100_000)}_{file_obj.filename}"

    try:
        # 4) Upload to S3
        s3.upload_fileobj(file_obj, S3_BUCKET, object_key)

        # 5) Generate presigned URL
        presigned_url = s3.generate_presigned_url(
            "get_object",
            Params={"Bucket": S3_BUCKET, "Key": object_key},
            ExpiresIn=expiration,
        )

    except Exception as e:
        # Log the exception in real code!
        return jsonify({"error": "S3 operation failed", "details": str(e)}), 500

    # 6) Return the URL
    return jsonify({"presigned_url": presigned_url}), 200
@app.route("/ask_question", methods=["POST"])
def ask_question():
    try:
        data = request.get_json()
        user_question = data.get("question", "")
        
        if not user_question:
            return jsonify({"error": "No question provided"}), 400

        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        
        # ✅ Fix: Enable safe deserialization
        vector_store = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
        
        docs = vector_store.similarity_search(user_question)

        chain = get_conversational_chain()
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)

        return jsonify({"response": response["output_text"]}), 200
    except Exception as e:
        import traceback
        traceback.print_exc()  # Prints full error in Flask console
        return jsonify({"error": str(e)}), 500
def mergePDF(pdf_list):
    merger = PdfWriter()
    for pdf in pdf_list:
        merger.append(pdf)
    output_path = tempfile.mktemp(suffix=".pdf")
    merger.write(output_path)
    merger.close()
    for i in pdf_list:
        print(i)
        os.remove(f'{i}')
    return output_path

def split_pdf(input_pdf, start_page, end_page):
    with open(input_pdf, 'rb') as input_file:
        reader = PyPDF2.PdfReader(input_file)
        writer = PyPDF2.PdfWriter()

        start_page = max(0, start_page - 1)
        end_page = min(len(reader.pages), end_page)

        for page_num in range(start_page, end_page):
            writer.add_page(reader.pages[page_num])
        output_path = tempfile.mktemp(suffix=".pdf")
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
    os.remove(input_pdf)
    return output_path
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
def summarise_commit(diff):
    response = client.chat.completions.create(
    model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": """You are an expert programmer, and you are trying to summarize a git diff.
    Reminders about the git diff format:
    For every file, there are a few metadata lines, like (for example):
    ```
    diff --git a/lib/index.js b/lib/index.js
    index aadf691..bfef603 100644
    --- a/lib/index.js
    +++ b/lib/index.js
    ```
    This means that `lib/index.js` was modified in this commit. Note that this is only an example.
    Then there is a specifier of the lines that were modified.
    A line starting with `+` means it was added.
    A line that starting with `-` means that line was deleted.
    A line that starts with neither `+` nor `-` is code given for context and better understanding.
    It is not part of the diff.
    [...]
    EXAMPLE SUMMARY COMMENTS:
    ```
    * Raised the amount of returned recordings from `10` to `100` [packages/server/recordings_api.ts], [packages/server/constants.ts]
    * Fixed a typo in the github action name [.github/workflows/gpt-commit-summarizer.yml]
    * Moved the `octokit` initialization to a separate file [src/octokit.ts], [src/index.ts]
    * Added an OpenAI API for completions [packages/utils/apis/openai.ts]
    * Lowered numeric tolerance for test files
    ```
    Most commits will have less comments than this examples list.
    The last comment does not include the file names,
    because there were more than two relevant files in the hypothetical commit.
    Do not include parts of the example in your summary.
    It is given only as an example of appropriate comments.""",
            },
            {
                "role": "user",
                "content": f"""Please summarise the following diff file: \n\n{diff}
                    
                    """,
            },
        ],
    )

    return response.choices[0].message.content.strip()
@app.route("/summarise-commit", methods=["POST"])
def summarise_commits():
    try:
        data = request.get_json()
        github_url = data.get("github_url")
        

        if not github_url:
            return jsonify({"error": "Missing github_url or commitHash"}), 400

        response = requests.get(
            f"{github_url}.diff",
            headers={
                "Accept": "application/vnd.github.v3.diff",
                "Authorization": f"token {os.getenv('GITHUB_TOKEN')}",
            },
        )

        if response.status_code != 200:
            return jsonify({"error": "Failed to fetch commit diff"}), response.status_code

        diff_content = response.text[:10000]  # Only pass first 10000 characters
        summary = summarise_commit(diff_content)

        return jsonify({"summary": summary})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
def pdfRotate(input_file, angle):
    reader = PdfReader(input_file)
    writer = PdfWriter()
    count = len(reader.pages)
    for i in range(count):
        writer.add_page(reader.pages[i])
        writer.pages[i].rotate(angle)

    output_path = tempfile.mktemp(suffix=".pdf")
    with open(output_path, "wb") as fp:
        writer.write(fp)
    os.remove(input_file)
    return output_path

def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    os.remove(pdf_file)
    return text

def save_text_to_docx(text, docx_file):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(docx_file)

def resize_image(input_path, output_path, target_width, target_height, max_size_kb=25):
    with Image.open(input_path) as img:
        original_width, original_height = img.size
        target_ratio = target_width / target_height
        original_ratio = original_width / original_height

        if target_ratio > original_ratio:
            # Target aspect ratio is wider than original
            new_height = int(target_width / original_ratio)
            img = img.resize((target_width, new_height), Image.Resampling.LANCZOS)
            crop_height = (new_height - target_height) // 2
            img = img.crop((0, crop_height, target_width, crop_height + target_height))
        else:
            # Target aspect ratio is taller than original
            new_width = int(target_height * original_ratio)
            img = img.resize((new_width, target_height), Image.Resampling.LANCZOS)
            crop_width = (new_width - target_width) // 2
            img = img.crop((crop_width, 0, crop_width + target_width, target_height))

        # Determine the format from the output file extension
        output_format = output_path.split('.')[-1].upper()
        if output_format == 'JPG':
            output_format = 'JPEG'
        quality = 95
        # Save the image with high quality settings initially
        if output_format in ['JPEG', 'JPG']:
            quality = 95
            img.save(output_path, format=output_format, quality=quality, optimize=True, progressive=True)
        else:
            img.save(output_path, format=output_format)
        
        file_size_kb = os.path.getsize(output_path) / 1024
        if file_size_kb > max_size_kb:
            quality -= 5
            while quality > 10:
                img.save(output_path, format=output_format, quality=quality, optimize=True, progressive=True)
                file_size_kb = os.path.getsize(output_path) / 1024
                if file_size_kb <= max_size_kb:
                    break
                quality -= 5
        
        print(f"Image resized and saved to {output_path} with size {file_size_kb:.2f} KB")
    os.remove(input_path)
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

@app.route('/')
def home():
    return jsonify({"message": "Excel-CSV Converter API is running!"})
@app.route('/pdf-to-docx', methods=['POST'])
def pdf_to_docx():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
            file.save(temp_pdf.name)
            pdf_path = temp_pdf.name

        output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
        try:
            cv = Converter(pdf_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            os.unlink(pdf_path)  # Delete temp PDF file

        return send_file(output_path, as_attachment=True, download_name="converted.docx")


@app.route('/docx-to-pdf', methods=['POST'])
def docx_to_pdf():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    try:
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")
        pdf_path = os.path.join(temp_dir, f"{uuid.uuid4()}.pdf")

        file.save(docx_path)
        convert(docx_path, pdf_path)

        if not os.path.exists(pdf_path):
            return jsonify({"error": "PDF conversion failed"}), 500

        @after_this_request
        def cleanup(response):
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                app.logger.warning(f"Cleanup failed: {e}")
            return response

        return send_file(pdf_path, as_attachment=True, download_name="converted.pdf")

    except Exception as e:
        return jsonify({"error": str(e)}), 500
@app.route("/generate-random", methods=["POST"])
def api_generate_random():
    data = request.json
    start = int(data.get("start", 10))
    end = int(data.get("end", 100))
    result = random.randrange(start,end)
    return jsonify({"random": result})
@app.route("/generate-uuids", methods=["POST"])
def api_generate_uuids():
    data = request.json
    num = int(data.get("num", 10))
    uuids = [str(uuid.uuid4()) for _ in range(num)]
    return jsonify({"uuids": uuids})
@app.route('/merge', methods=['POST'])
def merge():
    ensure_uploads_dir()
    files = request.files.getlist('files')
    file_paths = []
    for file in files:
        file_path = os.path.join('uploads', file.filename)
        file.save(file_path)
        file_paths.append(file_path)

    output_path = mergePDF(file_paths)

    @after_this_request
    def remove_files(response):
        try:
            os.remove(output_path)
            for file_path in file_paths:
                os.remove(file_path)
        except Exception as e:
            print(f"Error removing file: {e}")
        return response

    return send_file(output_path, as_attachment=True)
    
@app.route('/split', methods=['POST'])
def split():
    ensure_uploads_dir()
    file = request.files['file']
    start_page = int(request.form['start_page'])
    end_page = int(request.form['end_page'])

    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)

    output_path = split_pdf(file_path, start_page, end_page)

    @after_this_request
    def remove_files(response):
        try:
            os.remove(output_path)
            os.remove(file_path)
        except Exception as e:
            print(f"Error removing file: {e}")
        return response

    return send_file(output_path, as_attachment=True)
#hello
@app.route('/rotate', methods=['POST'])
def rotate():
    ensure_uploads_dir()
    file = request.files['file']
    angle = int(request.form['angle'])

    file_path = os.path.join('uploads', file.filename)
    file.save(file_path)

    output_path = pdfRotate(file_path, angle)

    @after_this_request
    def remove_files(response):
        try:
            os.remove(output_path)
            os.remove(file_path)
        except Exception as e:
            print(f"Error removing file: {e}")
        return response

    return send_file(output_path, as_attachment=True)
length_factors = {
    "mm": 0.001,
    "cm": 0.01,
    "m": 1,
    "km": 1000,
    "in": 0.0254,
    "ft": 0.3048,
    "yd": 0.9144,
    "mi": 1609.34
}

# Conversion factors for volume (to liters)
volume_factors = {
    "ml": 0.001,
    "l": 1,
    "m3": 1000,
    "gal": 3.78541,     # US gallon
    "qt": 0.946353,
    "pt": 0.473176,
    "cup": 0.24,
    "floz": 0.0295735
}

# Conversion factors for mass (to kilograms)
mass_factors = {
    "mg": 0.000001,
    "g": 0.001,
    "kg": 1,
    "ton": 1000,
    "lb": 0.453592,
    "oz": 0.0283495
}

@app.route('/convert_length', methods=['POST'])
def convert_length():
    data = request.get_json()
    value = data['value']
    from_unit = data['from_unit']
    to_unit = data['to_unit']

    meters = value * length_factors[from_unit]
    result = meters / length_factors[to_unit]
    return jsonify({
        "result": round(result, 4),
        "unit": to_unit
    })

@app.route('/convert_volume', methods=['POST'])
def convert_volume():
    data = request.get_json()
    value = data['value']
    from_unit = data['from_unit']
    to_unit = data['to_unit']

    liters = value * volume_factors[from_unit]
    result = liters / volume_factors[to_unit]
    return jsonify({
        "result": round(result, 4),
        "unit": to_unit
    })

@app.route('/convert_mass', methods=['POST'])
def convert_mass():
    data = request.get_json()
    value = data['value']
    from_unit = data['from_unit']
    to_unit = data['to_unit']

    kg = value * mass_factors[from_unit]
    result = kg / mass_factors[to_unit]
    return jsonify({
        "result": round(result, 4),
        "unit": to_unit
    })

@app.route('/convert_temperature', methods=['POST'])
def convert_temperature():
    data = request.get_json()
    value = data['value']
    from_unit = data['from_unit']
    to_unit = data['to_unit']

    def to_celsius(val, unit):
        if unit == "C": return val
        if unit == "F": return (val - 32) * 5 / 9
        if unit == "K": return val - 273.15

    def from_celsius(val, unit):
        if unit == "C": return val
        if unit == "F": return (val * 9 / 5) + 32
        if unit == "K": return val + 273.15

    celsius = to_celsius(value, from_unit)
    result = from_celsius(celsius, to_unit)
    return jsonify({
        "result": round(result, 2),
        "unit": to_unit
    })
@app.route('/convert_img', methods=['POST'])
def convert_image():
    if 'image' not in request.files:
        return jsonify({'error': 'No image file uploaded'}), 400
    if 'format' not in request.form:
        return jsonify({'error': 'Target format not specified'}), 400

    file = request.files['image']
    target_format = request.form['format'].lower()

    # Validate format
    valid_formats = ['jpeg', 'jpg', 'png', 'webp', 'bmp', 'gif', 'tiff']
    if target_format not in valid_formats:
        return jsonify({'error': f'Unsupported format: {target_format}'}), 400

    try:
        # Open the image
        img = Image.open(file.stream).convert("RGB")  # Convert to RGB to avoid mode issues

        # Prepare in-memory file
        img_io = io.BytesIO()
        img.save(img_io, target_format.upper())
        img_io.seek(0)

        # Determine MIME type
        mime_type = f'image/{target_format if target_format != "jpg" else "jpeg"}'

        return send_file(img_io, mimetype=mime_type, as_attachment=True,
                         download_name=f'converted.{target_format}')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
def extract_colors(img):
    tolerance = 32
    limit = 24
    colors, pixel_count = extcolors.extract_from_image(img, tolerance, limit)
    return colors

def render_color_platte(colors):
    size = 100
    columns = 6
    width = int(min(len(colors), columns) * size)
    height = int((math.floor(len(colors) / columns) + 1) * size)
    result = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    canvas = ImageDraw.Draw(result)
    for idx, color in enumerate(colors):
        x = int((idx % columns) * size)
        y = int(math.floor(idx / columns) * size)
        canvas.rectangle([(x, y), (x + size - 1, y + size - 1)], fill=color[0])
    return result

def overlay_palette(img, color_palette):
    f = plt.figure(figsize=(10, 8), facecolor='None', edgecolor='k', dpi=100)
    gs = gridspec.GridSpec(2, 1, wspace=0.0, hspace=0.0)
    ax1 = f.add_subplot(gs[0])
    ax1.imshow(img)
    ax1.axis('off')
    ax2 = f.add_subplot(gs[1])
    ax2.imshow(color_palette)
    ax2.axis('off')
    plt.subplots_adjust(wspace=0, hspace=0, bottom=0)

    buf = io.BytesIO()
    plt.savefig(buf, format='PNG', bbox_inches='tight', pad_inches=0)
    buf.seek(0)
    plt.close(f)
    return buf

def rgb_to_hex(rgb_tuple):
    return '#{:02x}{:02x}{:02x}'.format(*rgb_tuple)

@app.route('/extract-colors', methods=['POST'])
def extract_colors_api():
    if 'image' not in request.files:
        return jsonify({'error': 'No image uploaded'}), 400

    file = request.files['image']
    img = Image.open(file.stream).convert('RGB')
    print(img)
    colors = extract_colors(img)
    color_palette_img = render_color_platte(colors[:4])
    color_palette_data = io.BytesIO()
    color_palette_img.save(color_palette_data, format="PNG")
    color_palette_data.seek(0)
    palette_base64 = base64.b64encode(color_palette_data.read()).decode('utf-8')

    # Prepare color data with pixel count and hex
    hex_colors = [
        {
            'hex': rgb_to_hex(color[0]),
            'rgb': color[0],
            'pixels': color[1]
        }
        for color in colors
    ]

    return jsonify({
        'colors': hex_colors[:4],
        'palette_image_base64': palette_base64
    })
from flask import Flask, request, jsonify, send_file
import barcode
from barcode.writer import ImageWriter
import qrcode
import socket
import subprocess
import platform
import dns.resolver
import secrets
import string
import math



def generate_secure_password(length=20, use_symbols=True):
    if length < 8:
        raise ValueError("Password length should be at least 8 characters for good security.")
    alphabet = string.ascii_letters + string.digits
    if use_symbols:
        alphabet += string.punctuation

    while True:
        password = ''.join(secrets.choice(alphabet) for _ in range(length))
        if (any(c.islower() for c in password) and
            any(c.isupper() for c in password) and
            any(c.isdigit() for c in password) and
            (not use_symbols or any(c in string.punctuation for c in password))):
            return password

def calculate_entropy(password):
    charset_size = 0
    if any(c.islower() for c in password):
        charset_size += 26
    if any(c.isupper() for c in password):
        charset_size += 26
    if any(c.isdigit() for c in password):
        charset_size += 10
    if any(c in string.punctuation for c in password):
        charset_size += len(string.punctuation)
    if any(c.isspace() for c in password):
        charset_size += 1

    entropy = len(password) * math.log2(charset_size) if charset_size else 0
    return entropy, charset_size

def estimate_crack_time(entropy):
    guesses_per_second = 1e10
    total_combinations = 2 ** entropy
    seconds = total_combinations / guesses_per_second
    return convert_time(seconds)

def convert_time(seconds):
    time_units = [
        ("years", 60 * 60 * 24 * 365),
        ("months", 60 * 60 * 24 * 30),
        ("days", 60 * 60 * 24),
        ("hours", 60 * 60),
        ("minutes", 60),
        ("seconds", 1),
    ]
    result = []
    for unit, duration in time_units:
        if seconds >= duration:
            value = int(seconds // duration)
            seconds %= duration
            result.append(f"{value} {unit}")
    return ', '.join(result) if result else "less than a second"

def password_strength_feedback(entropy):
    if entropy < 28:
        return "Very Weak 🔴"
    elif entropy < 36:
        return "Weak 🟠"
    elif entropy < 60:
        return "Reasonable 🟡"
    elif entropy < 80:
        return "Strong 🟢"
    else:
        return "Very Strong 🔵"
currency = {'Afganistan': 'AFN', 'Albania': 'ALL', 'Alergia': 'DZD', 'American Samoa': 'USD', 'Andorra': 'EUR', 'Angola': 'AOA', 'Anguilla': 'XCD', 'Antigua and Barbuda': 'XCD', 'Argentina': 'ARS', 'Armenia': 'AMD', 'Aruba': 'AWG', 'Australia': 'AUD', 'Austria': 'EUR', 'Azerbaijan': 'AZN', 'Bahamas': 'BSD', 'Bahrain': 'BHD', 'Bangladesh': 'BDT', 'Barbados': 'BBD', 'Belarus': 'BYN', 'Belgium': 'EUR', 'Belize': 'BZD', 'Benin': 'CFA', 'Bermuda': 'BMD', 'Bhutan': 'BTN', 'Bolivia': 'BOV', 'Bonaire': 'BES', 'Boznia and herzegovina': 'BAM', 'Botswana': 'BWP', 'Bouvet': 'NOK', 'Brazil': 'BRL', 'British Indian Ocean Terriotory': 'USD', 'Brunei Daraussalam': 'BND', 'Bulgaria': 'BGN', 'Burkina': 'CFA', 'Burundi': 'BIF', 'Cape Verde': 'CVE', 'Cambodia': 'KHR', 'Cameroon': 'CFA', 'Canada': 'CAD', 'Cayman Islands': 'KYD', 'Central African Repulic': 'CFA', 'Chad': 'CFA', 'Chile': 'CLP', 'China': 'CNY', 'Christmas Island': 'AUD', 'Cocos Islands': 'AUD', 'Colombia': 'COP', 'Comoros': 'KMF', 'The Democratic Repulic of Congo': 'CDF', 'The Cook Islands': 'NZD', 'Costa Rica': 'CRC', 'Croatia': 'HRK', 'Cuba': 'CUP', 'CuraÇao': 'ANG', 'Cyprus': 'EUR', 'Czech Repulic': 'CZK', 'Ivory Coast': 'CFA', 'Denmark': 'DKK', 'Djibouti': 'DJF', 'Dominica': 'XCD', 'The Dominican Republic': 'DOP', 'Equador': 'USD', 'Egypt': 'EGP', 'El Salvador': 'USD', 'Equatorial Guniea': 'CFA', 'Eritrea': 'ERN', 'Estonia': 'EUR', 'Ethipia': 'ETB', 'The Falkland': 'FKP', 'The Faroe': 'DKK', 'Fiji': 'FJD', 'Finland': 'EUR', 'France': 'EUR', 'French Guiana': 'EUR', 'French Polynesia': 'CFP', 'French Southern Territores': 'EUR', 'Gabon': 'CFA', 'Gambia': 'GMD', 'Georgia': 'GEL', 'Germany': 'EUR', 'Ghana': 'GHS', 'Gibraltar': 'GIP', 'Greece': 'EUR', 'Greenland': 'DKK', 'Grenada': 'XCD', 'Guadeloupe': 'EUR', 'Guam': 'USD', 'Guatemala': 'GTQ', 'Guernsey': 'GBP', 'Guniea': 'GNF', 'Guinea-Bisaau': 'CFA', 'Guyana': 'GYD', 'Haiti': 'USD', 'Holy See': 'EUR', 'Honduras': 'HNL', 'Hong Kong': 'HKD', 'Hungary': 'HUF', 'Iceland': 'ISK', 'India': 'INR', 'Indonesia': 'IDR', 'Iran': 'IRR', 'Iraq': 'IQD', 'Ireland': 'EUR', 'Isle of man': 'GBP', 'Israel': 'ILS', 'Italy': 'EUR', 'Jamaica': 'JMD', 'Japan': 'JPY', 'Jersey': 'GBP', 'Jordan': 'JOD', 'Kazakstan': 'KZT', 'Kenya': 'KES', 'Kiribati': 'AUD', 'North Korea': 'KPW', 'South Korea': 'KRW', 'Kuwait': 'KWD', 'Kyrgyzstan': 'KGS', 'Lao': 'LAK', 'Latvia': 'EUR', 'Lebanon': 'LBP', 'Lesotho': 'ZAR', 'Liberia': 'LRD', 'Libya': 'LYD', 'Liechensteain': 'CHF', 'Lithuania': 'EUR', 'Luxembourg': 'EUR', 'Macao': 'MOP', 'Madagascar': 'MGA', 'Malawi': 'MWK', 'Malaysia': 'MYR', 'Maldives': 'MVR', 'Mali': 'CFA', 'Malta': 'EUR', 'The Marshall Islands': 'USD', 'Martinque': 'EUR', 'Mauritania': 'MRU', 'Mauritius': 'MUR', 'Mayotte': 'EUR', 'Mexico': 'MXV', 'Micronesia': 'USD', 'Moldova': 'MDL', 'Monaco': 'EUR', 'Mongolia': 'MNT', 'Montenegro': 'EUR', 'Montserrat': 'XCD', 'Morocco': 'MAD', 'Mozambique': 'MZN', 'Myanmar': 'MMK', 'Namibia': 'ZAR', 'Nauru': 'AUD', 'Nepal': 'NPR', 'The Netherlands': 'EUR', 'New Caledonia': 'CFP', 'New Zealand': 'NZD', 'Nicaragua': 'NIO', 'Niger': 'CFA', 'Nigeria': 'NGN', 'Niue': 'NZD', 'Nolfolf Island': 'AUD', 'Northern Mariana Islands': 'USD', 'Norway': 'NOK', 'Oman': 'OMR', 'Pakistan': 'PKR', 'Palau': 'USD', 'Panama': 'USD', 'Papua New Guinea': 'PGK', 'Paraguay': 'PYG', 'Peru': 'PEN', 'Philippines': 'PHP', 'Pitcairn': 'NZD', 'Poland': 'PLN', 'Portugal': 'EUR', 'Puerto Rico': 'USD', 'Qatar': 'QAR', 'North Macedonia': 'MKD', 'Romania': 'RON', 'Russia': 'RUB', 'Rwanda': 'RWF', 'Réunion': 'EUR', 'Saint Barts': 'EUR', 'Saint Helena': 'SHP', 'Saint Kitts and Nevis': 'XCD', 'Saint Lucia': 'XCD', 'Saint Martin': 'EUR', 'Saint Pierre and Miquelon': 'EUR', 'Saint Vincent and the Grenadines': 'XCD', 'Samoa': 'WST', 'San Marino': 'EUR', 'Sao Tome and Principe': 'STN', 'Saudi Arabia': 'SAR', 'Senegal': 'CFA', 'Serbia': 'RSD', 'Seychelles': 'SCR', 'Sierra Leone': 'SLL', 'Singapore': 'SGD', 'Sint Maarten': 'ANG', 'Slovakia': 'EUR', 'Slovenia': 'EUR', 'Solomon Islands': 'SBD', 'Somalia': 'SOS', 'South Africa': 'ZAR', 'South Sudan': 'SSP', 'Spain': 'EUR', 'Sri Lanka': 'LKR', 'Sudan': 'SDG', 'Suriname': 'SRD', 'Svalbard and Jan Mayen': 'NOK', 'Swaziland': 'SZL', 'Sweden': 'SEK', 'Switzerland': 'CHF', 'Syria': 'SYP', 'Taiwan': 'TWD', 'Tajikistan': 'TJS', 'Tanzania': 'TZS', 'Thailand': 'THB', 'Timor-leste': 'USD', 'Togo': 'CFA', 'Tokelau': 'NZD', 'Tonga': 'TOP', 'Trinidad and Tobago': 'TTD', 'Tunisia': 'TND', 'Turkey': 'TRY', 'Turkmenistan': 'TMT', 'Turks and Caicos Islands': 'USD', 'Tuvalu': 'AUD', 'Uganda': 'UGX', 'Ukraine': 'UAH', 'United Arab Emirates': 'UAE', 'United Kingdom': 'GBP', 'United States Minor Outlying Islands': 'USD', 'United States of America': 'USD', 'Uruguay': 'UYU', 'Uzbekistan': 'UZS', 'Vanuatu': 'VUV', 'Venezuela': 'VEF', 'Vietnam': 'VND', 'British Virgin Islands': 'USD', 'US Virgin Islands': 'USD', 'Wallis and Futuna': 'CFP', 'Western Sahara': 'MAD', 'Yemen': 'YER', 'Zambia': 'ZMW', 'Zimbabwe': 'ZWL', 'Åland Islands': 'EUR'}

# ------------------- Routes -------------------

@app.route("/generate-password", methods=["POST"])
def api_generate_password():
    length = int(request.json.get("length", 20))
    use_symbols = "true"
    try:
        password = generate_secure_password(length, use_symbols)
        return jsonify({"password": password})
    except ValueError as e:
        return jsonify({"error": str(e)}), 400

@app.route("/check-password", methods=["POST"])
def api_check_password():
    data = request.json
    password = data.get("password", "")
    if not password:
        return jsonify({"error": "No password provided"}), 400

    entropy, charset = calculate_entropy(password)
    crack_time = estimate_crack_time(entropy)
    strength = password_strength_feedback(entropy)

    return jsonify({
        "password": password,
        "charset_size": charset,
        "entropy": round(entropy, 2),
        "crack_time": crack_time,
        "strength_feedback": strength
    })

@app.route("/ping", methods=["POST"])
def api_ping():
    host = request.json.get("host")
    print(host)
    if not host:
        return jsonify({"error": "No host provided"}), 400
    param = "-n" if platform.system().lower() == "windows" else "-c"
    try:
        output = subprocess.check_output(["ping", param, "4", host], stderr=subprocess.STDOUT, text=True)
        return jsonify({"result": output})
    except subprocess.CalledProcessError as e:
        return jsonify({"error": e.output}), 500

@app.route("/dns-lookup", methods=["POST"])
def api_dns_lookup():
    domain = request.json.get("host")
    if not domain:
        return jsonify({"error": "No domain provided"}), 400
    try:
        result = dns.resolver.resolve(domain, 'A')
        return jsonify({"dns_records": [ip.address for ip in result]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/ip-lookup", methods=["POST"])
def api_ip_lookup():
    domain = request.json.get("host")
    if not domain:
        return jsonify({"error": "No domain provided"}), 400
    try:
        ip = socket.gethostbyaddr(domain)
        return jsonify({"ip": ip})
    except socket.gaierror as e:
        return jsonify({"error": str(e)}), 500

@app.route("/traceroute", methods=["POST"])
def api_traceroute():
    host = request.json.get("host")
    if not host:
        return jsonify({"error": "No host provided"}), 400
    traceroute_cmd = "tracert" if platform.system().lower() == "windows" else "traceroute"
    try:
        output = subprocess.check_output([traceroute_cmd, host], stderr=subprocess.STDOUT, text=True)
        return jsonify({"result": output})
    except subprocess.CalledProcessError as e:
        return jsonify({"error": e.output}), 500
@app.route('/convert_currency', methods=['POST'])
def currency_convert():
    data = request.json
    country_from = data.get('from')
    country_to = data.get('to')
    amount = data.get('amount')

    if not all([country_from, country_to, amount]):
        return jsonify({'error': 'Missing required fields: from, to, amount'}), 400

    try:
        amount = float(amount)
        from_code = currency.get(country_from)
        to_code = currency.get(country_to)

        if not from_code or not to_code:
            return jsonify({'error': 'Invalid country name'}), 400

        symbols = f"{from_code},{to_code}"
        url = f"https://api.currencyfreaks.com/v2.0/rates/latest?apikey={api_key}&symbols={symbols}"

        response = requests.get(url)
        data = response.json()
        print(data)
        rate_from = float(data['rates'][from_code])
        rate_to = float(data['rates'][to_code])

        final_amount = amount * rate_to / rate_from
        print(final_amount)
        return jsonify({
            'from': country_from,
            'to': country_to,
            'amount': amount,
            'converted': round(final_amount, 2),
            'from_code': from_code,
            'to_code': to_code
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500
@app.route("/generate-barcode", methods=["POST"])
def api_generate_barcode():
    # Get text from form-data
    text = request.form.get("text", "sample")

    try:
        barcode_format = barcode.get_barcode_class('code128')
        barcode_obj = barcode_format(text, writer=ImageWriter())
        filename = "barcode_output"
        barcode_path = barcode_obj.save(filename)

        return send_file(f"D:\\Codeshastra_XI_SleepyHeads\\barcode_output.png", mimetype='image/png')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/generate-qrcode", methods=["POST"])
def api_generate_qrcode():
    # Get text from form-data
    text = request.form.get("text")
    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        img = qrcode.make(text)
        qr_path = "ml/qrcode.png"
        img.save(qr_path)

        return send_file('qrcode.png', mimetype='image/png')
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/excel-to-csv', methods=['POST'])
def excel_to_csv_endpoint():
    # Check if file is in the request
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file and file.filename.endswith(('.xlsx', '.xls')):
        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(excel_path)
        
        # Create output path
        csv_filename = os.path.splitext(filename)[0] + '.csv'
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], csv_filename)
        
        try:
            # Convert Excel to CSV
            df = pd.read_excel(excel_path, engine='openpyxl')
            df.to_csv(csv_path, index=False)
            
            # Return the CSV file
            return send_file(
                csv_path,
                mimetype='text/csv',
                as_attachment=True,
                download_name=csv_filename
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            # Clean up files
            if os.path.exists(excel_path):
                os.remove(excel_path)
            # Keep CSV file for a short time for download, it will be removed later
    else:
        return jsonify({"error": "File must be an Excel file (.xlsx or .xls)"}), 400

@app.route('/csv-to-excel', methods=['POST'])
def csv_to_excel_endpoint():
    # Check if file is in the request
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file and file.filename.endswith('.csv'):
        # Save uploaded file temporarily
        filename = secure_filename(file.filename)
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(csv_path)
        
        # Create output path
        excel_filename = os.path.splitext(filename)[0] + '.xlsx'
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], excel_filename)
        
        try:
            # Convert CSV to Excel
            df = pd.read_csv(csv_path)
            df.to_excel(excel_path, index=False, engine='openpyxl')
            
            # Return the Excel file
            return send_file(
                excel_path,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name=excel_filename
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            # Clean up files
            if os.path.exists(csv_path):
                os.remove(csv_path)
            # Keep Excel file for a short time for download, it will be removed later
    else:
        return jsonify({"error": "File must be a CSV file (.csv)"}), 400

# Cleanup task to remove old files (in a production environment, you might use a scheduled task)
@app.route('/cleanup', methods=['POST'])
def cleanup():
    count = 0
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
                count += 1
        except Exception as e:
            print(f"Error deleting {file_path}: {e}")
    
    return jsonify({"message": f"Removed {count} files"})

if __name__ == '__main__':
    app.run(debug=True)